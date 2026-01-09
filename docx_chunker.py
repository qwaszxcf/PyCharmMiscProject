import json
import os
from typing import List, Dict
from datetime import datetime

from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LCDocument



def load_docx_sections(file_path: str) -> List[Dict]:
    """
    读取 docx，将内容按标题（Heading）进行分组，返回一个 section 列表。
    每个 section: {title, level, content}
    """
    doc = Document(file_path)

    sections = []
    current_title = None
    current_level = None
    current_paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style is not None else ""

        if style_name.startswith("Heading"):
            # 先把上一个 section 收尾
            if current_paragraphs:
                sections.append(
                    {
                        "title": current_title,
                        "level": current_level,
                        "content": "\\n".join(current_paragraphs),
                    }
                )
                current_paragraphs = []

            current_title = text
            # 尝试从样式名中提取层级数字
            try:
                level = int("".join([c for c in style_name if c.isdigit()]) or "1")
            except ValueError:
                level = 1
            current_level = level
        else:
            # 普通段落，归入当前 section（如果没有标题，可以归为"文档开头"）
            if current_title is None:
                current_title = "文档开头"
                current_level = 0
            current_paragraphs.append(text)

    # 处理最后一个 section
    if current_paragraphs:
        sections.append(
            {
                "title": current_title,
                "level": current_level,
                "content": "\\n".join(current_paragraphs),
            }
        )

    return sections


def sections_to_documents(sections: List[Dict]) -> List[LCDocument]:
    """把 section 转换为 LangChain 的 Document 对象，带元数据。"""
    docs = []
    for sec in sections:
        metadata = {
            "title": sec["title"],
            "level": sec["level"],
        }
        docs.append(LCDocument(page_content=sec["content"], metadata=metadata))
    return docs


def chunk_documents_with_recursive_splitter(
    docs: List[LCDocument],
    chunk_size: int = 500,
    chunk_overlap: int = 100,
    source: str = "",
) -> List[LCDocument]:
    """
    使用 RecursiveCharacterTextSplitter 按字符数切分。
    标题不切分，只对正文内容进行切分，每个chunk保留原标题信息。
    chunk_size: 每块最大字符数
    chunk_overlap: 相邻块的重叠字符数
    source: 源文档路径
    """
    # 按照层级顺序进行分割，优先级为：\n\n\n, \n\n, \n, " ", ""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n\n", "\n\n", "\n", " ", ""],
    )
    
    # 对每个文档进行切分，并添加详细的metadata
    all_chunked_docs = []
    global_chunk_index = 0
    
    for doc in docs:
        # 获取原始标题和层级信息
        original_title = doc.metadata.get("title", "未知标题")
        original_level = doc.metadata.get("level", 0)
        
        # 只对正文内容进行切分（标题不参与切分）
        content = doc.page_content
        
        # 如果内容为空，跳过
        if not content or not content.strip():
            continue
        
        # 使用splitter对正文内容进行切分
        text_chunks = splitter.split_text(content)
        
        # 为每个切分后的文本块创建Document，并保留原标题
        for chunk_text in text_chunks:
            chunk_doc = LCDocument(
                page_content=chunk_text,
                metadata={
                    "title": original_title,  # 保持原标题不变
                    "level": original_level,  # 保持原层级不变
                    "chunk_index": global_chunk_index,
                    "source": source,
                }
            )
            all_chunked_docs.append(chunk_doc)
            global_chunk_index += 1
    
    return all_chunked_docs


def process_docx_for_rag(file_path: str, chunk_size: int = 500, chunk_overlap: int = 50) -> List[LCDocument]:
    """
    处理 Word 文档并返回切分后的文档块
    """
    print(f"正在处理文档: {file_path}")
    
    # 1. 读取 docx -> sections
    sections = load_docx_sections(file_path)
    print(f"文档解析完成，共 {len(sections)} 个章节")
    
    # 2. sections -> LangChain Document
    base_docs = sections_to_documents(sections)
    
    # 3. 按字符递归分块
    chunked_docs = chunk_documents_with_recursive_splitter(
        base_docs,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        source=file_path,
    )
    
    print(f"文档切分完成，共生成 {len(chunked_docs)} 个 chunk")
    
    return chunked_docs


def convert_docs_to_chunks(
    docs: List[LCDocument],
    doc_id_prefix: str = "doc"
) -> List[Dict]:
    """
    将 LangChain Document 列表转换为标准 chunk 格式
    
    Args:
        docs: LangChain Document 列表
        doc_id_prefix: chunk_id 前缀
        
    Returns:
        chunk 列表，格式:
        [
            {
                "chunk_id": "doc_chunk_0001",
                "text": "chunk 文本内容",
                "metadata": {"title": "...", "level": 1, ...}
            },
            ...
        ]
    """
    chunks = []
    for i, doc in enumerate(docs):
        chunk_id = f"{doc_id_prefix}_chunk_{i:04d}"
        chunks.append({
            "chunk_id": chunk_id,
            "text": doc.page_content,
            "metadata": dict(doc.metadata) if hasattr(doc, 'metadata') else {}
        })
    return chunks


def save_chunks_to_json(
    chunks: List[Dict],
    output_path: str,
    source_file: str = None
) -> str:
    """
    将 chunk 列表保存为 JSON 文件（冻结 chunk 结构）
    
    Args:
        chunks: chunk 列表
        output_path: 输出文件路径
        source_file: 源文档路径（可选，用于记录来源）
        
    Returns:
        输出文件路径
    """
    output_data = {
        "metadata": {
            "source_file": source_file,
            "total_chunks": len(chunks),
            "created_at": datetime.now().isoformat()
        },
        "chunks": chunks
    }
    
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(output_data, f, ensure_ascii=False, indent=2)
    
    print(f"已保存 {len(chunks)} 个 chunk 到: {output_path}")
    return output_path


def load_chunks_from_json(input_path: str) -> List[Dict]:
    """
    从 JSON 文件加载 chunk 列表
    
    Args:
        input_path: JSON 文件路径
        
    Returns:
        chunk 列表
    """
    with open(input_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    chunks = data.get("chunks", [])
    print(f"已加载 {len(chunks)} 个 chunk 从: {input_path}")
    return chunks


def process_and_save_chunks(
    file_path: str,
    output_path: str = None,
    doc_id_prefix: str = None,
    chunk_size: int = 500,
    chunk_overlap: int = 50
) -> List[Dict]:
    """
    处理文档并保存 chunk 到 JSON 文件（一站式接口）
    
    Args:
        file_path: docx 文件路径
        output_path: 输出 JSON 文件路径（默认为 {文件名}_chunks.json）
        doc_id_prefix: chunk_id 前缀（默认为文件名）
        chunk_size: 每块最大字符数
        chunk_overlap: 相邻块的重叠字符数
        
    Returns:
        chunk 列表
    """
    # 确定输出路径和前缀
    base_name = os.path.splitext(os.path.basename(file_path))[0]
    if output_path is None:
        output_path = f"{base_name}_chunks.json"
    if doc_id_prefix is None:
        doc_id_prefix = base_name
    
    # 1. 切分文档
    docs = process_docx_for_rag(file_path, chunk_size, chunk_overlap)
    
    # 2. 转换为标准 chunk 格式
    chunks = convert_docs_to_chunks(docs, doc_id_prefix)
    
    # 3. 保存到 JSON
    save_chunks_to_json(chunks, output_path, source_file=file_path)
    
    return chunks


if __name__ == "__main__":
    # 处理 testdoc.docx 文件并保存 chunk JSON
    try:
        # 使用一站式接口：切分 + 转换 + 保存
        chunks = process_and_save_chunks(
            file_path="testdoc.docx",
            output_path="testdoc_chunks.json",
            chunk_size=500,
            chunk_overlap=50
        )
        
        print(f"\n共生成 {len(chunks)} 个 chunk")
        print("\n" + "=" * 50)
        print("Chunk 结构示例（前 3 个）:")
        print("=" * 50)
        
        for chunk in chunks[:3]:
            print(json.dumps(chunk, ensure_ascii=False, indent=2))
            print("-" * 40)
            
    except FileNotFoundError:
        print("错误: 找不到 testdoc.docx 文件，请确保该文件存在于当前目录中")
    except Exception as e:
        print(f"处理文档时发生错误: {str(e)}")