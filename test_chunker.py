"""
测试文档切割功能的脚本
此脚本将在 testdoc.docx 存在时运行文档切割功能
"""
import os
from docx_chunker import process_docx_for_rag

def main():
    # 检查 testdoc.docx 是否存在
    if not os.path.exists("testdoc.docx"):
        print("警告: testdoc.docx 文件不存在，将创建一个示例文档用于测试")
        
        # 创建一个示例 Word 文档用于测试
        from docx import Document
        
        doc = Document()
        
        # 添加标题和内容
        doc.add_heading('员工请假审批规则', 0)
        doc.add_paragraph('这是关于员工请假审批的基本规则文档。')
        
        doc.add_heading('请假类型', level=1)
        doc.add_paragraph('1. 病假：需要医院证明，最长可请30天。')
        doc.add_paragraph('2. 年假：根据工龄计算，最多可累积5天。')
        doc.add_paragraph('3. 事假：需提前3天申请，最长不超过7天。')
        
        doc.add_heading('审批流程', level=1)
        doc.add_paragraph('1. 员工提交请假申请')
        doc.add_paragraph('2. 直接主管审批')
        doc.add_paragraph('3. HR部门备案')
        
        doc.add_heading('特殊情况处理', level=1)
        doc.add_paragraph('1. 紧急情况可事后补交申请')
        doc.add_paragraph('2. 跨年请假需分段申请')
        
        doc.save('testdoc.docx')
        print("已创建示例文档 testdoc.docx")
    
    print("开始处理文档...")
    
    # 调用文档切割功能
    try:
        chunks = process_docx_for_rag("testdoc.docx", chunk_size=200, chunk_overlap=20)
        
        print(f"\n处理完成！共生成 {len(chunks)} 个文档块")
        print("\n前3个文档块预览:")
        for i, chunk in enumerate(chunks[:3]):
            print(f"\n--- 文档块 {i+1} ---")
            print(f"标题: {chunk.metadata.get('title', 'N/A')}")
            print(f"内容预览: {chunk.page_content[:100]}...")
            print(f"内容长度: {len(chunk.page_content)} 字符")
    
    except ImportError as e:
        print(f"缺少必要的依赖包: {e}")
        print("请运行以下命令安装依赖:")
        print("pip install python-docx langchain tiktoken")
    except Exception as e:
        print(f"处理文档时发生错误: {e}")

if __name__ == "__main__":
    main()